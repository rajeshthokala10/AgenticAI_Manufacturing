"""Main pipeline: orchestrates extraction, cleaning, chunking, and indexing."""

import logging
from pathlib import Path
from typing import List, Optional

from .models import (
    ProcessedChunk, PageInfo, PageType, DocumentType, ExtractionResult
)
from .config import PipelineConfig
from . import extractors, cleaners, chunkers, utils

logger = logging.getLogger(__name__)


class ProductionRAGPipeline:
    """
    Production RAG document processing pipeline.

    Handles 14+ data types and failure scenarios:
    1.  Tables (structured, nested, cross-page)
    2.  Headers, footers, page numbers
    3.  Multi-column layouts
    4.  Nested lists & numbered clauses
    5.  Scanned pages mixed with digital
    6.  Forms & key-value pairs
    7.  Watermarks, stamps & annotations
    8.  Embedded charts & graphs
    9.  Cross-references & footnotes
    10. Merged/concatenated documents
    11. Confidential/redacted content
    12. Metadata & document properties
    13. Encoding & unicode issues
    14. Password-protected PDFs
    15. Embedded file attachments
    16. TOC/index page filtering
    17. Chunk quality validation
    18. Document deduplication
    19. Document versioning
    20. Fault tolerance (per-page error handling)
    """

    def __init__(self, config: Optional[PipelineConfig] = None):
        self.config = config or PipelineConfig()

    def process(self, file_path: str, password: str = "") -> List[ProcessedChunk]:
        """Main entry point: process any supported document."""
        file_path = str(file_path)
        doc_type = self._detect_type(file_path)
        logger.info(f"Processing {Path(file_path).name} as {doc_type.value}")

        if doc_type == DocumentType.PDF:
            chunks = self._process_pdf(file_path, password)
        elif doc_type == DocumentType.DOCX:
            chunks = self._process_docx(file_path)
        elif doc_type == DocumentType.HTML:
            chunks = self._process_html(file_path)
        else:
            chunks = self._process_text(file_path)

        # Post-processing
        chunks = chunkers.validate_chunk_quality(chunks)
        chunks = utils.deduplicate_chunks(chunks, self.config)
        utils.generate_embeddings(chunks, self.config)

        # Versioning
        if self.config.enable_versioning:
            fp = utils.compute_doc_fingerprint(chunks)
            version_info = utils.check_version_changes(fp, file_path, self.config)
            if version_info:
                for c in chunks:
                    c.metadata["doc_version"] = version_info

        logger.info(f"Produced {len(chunks)} chunks from {Path(file_path).name}")
        return chunks

    # =========================================================================
    # PDF processing
    # =========================================================================

    def _process_pdf(self, file_path: str, password: str) -> List[ProcessedChunk]:
        chunks = []
        warnings = []

        # Step 1: Open (handles encryption)
        doc = utils.open_pdf_safe(file_path, password)
        total_pages = len(doc)

        # Step 2: Metadata
        metadata = extractors.extract_metadata(doc, file_path)

        # Step 3: Classify pages
        classification = extractors.classify_pages(doc, self.config)

        # Step 4: Detect headers/footers
        headers, footers = cleaners.detect_headers_footers(doc, self.config)

        # Step 5: Detect redactions
        redactions = {}
        if self.config.detect_redactions:
            redactions = cleaners.detect_redactions(doc)

        # Step 6: Detect document boundaries (merged PDFs)
        boundaries = utils.detect_document_boundaries(doc)

        # Step 7: Extract text page by page (with fault tolerance)
        pages = []
        max_pages = self.config.max_pages or total_pages
        for page_num in range(min(total_pages, max_pages)):
            try:
                page = doc[page_num]
                page_type = classification.get(page_num, PageType.DIGITAL)
                confidence = 1.0

                if page_type == PageType.DIGITAL:
                    text = extractors.extract_text_digital(page, self.config)
                elif page_type == PageType.SCANNED:
                    text, confidence = extractors.extract_text_ocr(
                        file_path, page_num, self.config
                    )
                else:
                    text = ""

                # Clean: strip headers/footers
                text = cleaners.strip_headers_footers(text, headers, footers)

                # Clean: remove watermarks
                if self.config.strip_watermarks:
                    watermarks = cleaners.detect_watermarks(page)
                    text = cleaners.remove_watermarks(text, watermarks)

                # Clean: mark redactions
                has_redactions = False
                if page_num in redactions:
                    text, has_redactions = cleaners.apply_redaction_markers(
                        page, redactions[page_num], text
                    )

                # Flag TOC pages
                is_toc = cleaners.is_toc_page(text) if self.config.filter_toc_pages else False

                pages.append(PageInfo(
                    page_num=page_num + 1, text=text, page_type=page_type,
                    confidence=confidence, has_redactions=has_redactions, is_toc=is_toc,
                ))

            except Exception as e:
                msg = f"Error on page {page_num + 1}: {e}"
                logger.error(msg)
                warnings.append(msg)
                if self.config.fail_on_error:
                    raise
                pages.append(PageInfo(
                    page_num=page_num + 1, text="", page_type=PageType.EMPTY,
                ))

        # Step 8: Extract tables
        try:
            tables = extractors.extract_tables(file_path, self.config)
            table_chunks = chunkers.chunk_tables(tables, self.config)
            chunks.extend(table_chunks)
        except Exception as e:
            logger.error(f"Table extraction failed: {e}")
            if self.config.fail_on_error:
                raise

        # Step 9: Extract forms
        try:
            forms = extractors.extract_forms(doc, self.config)
            form_chunks = chunkers.chunk_forms(forms, Path(file_path).name)
            chunks.extend(form_chunks)
        except Exception as e:
            logger.error(f"Form extraction failed: {e}")

        # Step 10: Extract charts
        if self.config.enable_vlm_charts:
            try:
                charts = extractors.extract_charts(doc, self.config)
                for chart in charts:
                    chart.description = extractors.describe_chart_with_vlm(chart, self.config)
                    chunks.append(ProcessedChunk(
                        content=f"[Chart, page {chart.page}] {chart.description}",
                        content_type="chart",
                        metadata={"page": chart.page, "dimensions": f"{chart.width}x{chart.height}"},
                    ))
            except Exception as e:
                logger.error(f"Chart extraction failed: {e}")

        # Step 11: Extract annotations
        try:
            annotations = cleaners.extract_annotations(doc)
            for ann in annotations:
                chunks.append(ProcessedChunk(
                    content=f"[{ann['type']} annotation, page {ann['page']}] {ann['content']}",
                    content_type="footnote",
                    metadata={"page": ann["page"], "author": ann.get("author", "")},
                ))
        except Exception as e:
            logger.error(f"Annotation extraction failed: {e}")

        # Step 12: Extract embedded attachments
        try:
            attachments = extractors.extract_attachments(doc)
            for att in attachments:
                chunks.append(ProcessedChunk(
                    content=f"[Embedded attachment: {att['name']}, {att['size']} bytes]",
                    content_type="text",
                    metadata={"attachment_name": att["name"], "attachment_size": att["size"]},
                ))
        except Exception as e:
            logger.error(f"Attachment extraction failed: {e}")

        doc.close()

        # Step 13: Parse sections & footnotes from full text
        content_pages = [p for p in pages if not p.is_toc]
        full_text = "\n\n".join(p.text for p in content_pages if p.text)
        sections = chunkers.parse_sections(full_text)
        footnotes = chunkers.extract_footnotes(full_text)
        section_lookup = {s.number: s.content[:300] for s in sections}

        # Step 14: Chunk text
        if sections:
            text_chunks = chunkers.chunk_with_hierarchy(sections, self.config)
        else:
            text_chunks = chunkers.semantic_chunk(
                [{"page": p.page_num, "text": p.text} for p in content_pages],
                self.config,
            )

        # Step 15: Resolve cross-references
        text_chunks = chunkers.resolve_references(
            text_chunks, footnotes, section_lookup, self.config
        )

        # Convert to ProcessedChunk
        for tc in text_chunks:
            chunks.append(ProcessedChunk(
                content=tc["content"],
                content_type=tc.get("metadata", {}).get("type", "text"),
                metadata={**tc.get("metadata", {}), **metadata,
                          "warnings": warnings if warnings else None},
            ))

        return chunks

    # =========================================================================
    # DOCX processing
    # =========================================================================

    def _process_docx(self, file_path: str) -> List[ProcessedChunk]:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)

        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extract tables from DOCX
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            if rows:
                headers = rows[0]
                md = chunkers._rows_to_md(headers, rows[1:])
                full_text.append(f"\n{md}\n")

        text = "\n".join(full_text)
        sections = chunkers.parse_sections(text)

        if sections:
            text_chunks = chunkers.chunk_with_hierarchy(sections, self.config)
        else:
            text_chunks = chunkers.semantic_chunk([{"page": 1, "text": text}], self.config)

        return [
            ProcessedChunk(
                content=tc["content"],
                content_type=tc.get("metadata", {}).get("type", "text"),
                metadata=tc.get("metadata", {}),
            )
            for tc in text_chunks
        ]

    # =========================================================================
    # HTML processing
    # =========================================================================

    def _process_html(self, file_path: str) -> List[ProcessedChunk]:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            html = f.read()

        try:
            import trafilatura
            text = trafilatura.extract(html, include_tables=True, include_links=True)
        except ImportError:
            from html.parser import HTMLParser

            class _Strip(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.parts = []
                def handle_data(self, d):
                    self.parts.append(d)

            s = _Strip()
            s.feed(html)
            text = "\n".join(s.parts)

        text_chunks = chunkers.semantic_chunk([{"page": 1, "text": text or ""}], self.config)
        return [
            ProcessedChunk(content=tc["content"], content_type="text", metadata=tc.get("metadata", {}))
            for tc in text_chunks
        ]

    # =========================================================================
    # Plain text processing
    # =========================================================================

    def _process_text(self, file_path: str) -> List[ProcessedChunk]:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

        text_chunks = chunkers.semantic_chunk([{"page": 1, "text": text}], self.config)
        return [
            ProcessedChunk(content=tc["content"], content_type="text", metadata=tc.get("metadata", {}))
            for tc in text_chunks
        ]

    # =========================================================================
    # File type detection
    # =========================================================================

    @staticmethod
    def _detect_type(file_path: str) -> DocumentType:
        ext = Path(file_path).suffix.lower()
        return {
            ".pdf": DocumentType.PDF, ".docx": DocumentType.DOCX,
            ".doc": DocumentType.DOCX, ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
        }.get(ext, DocumentType.TXT)
